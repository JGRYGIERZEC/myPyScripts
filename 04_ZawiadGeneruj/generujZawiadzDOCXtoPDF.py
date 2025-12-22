import shlex
import os
import sys
from collections import defaultdict
from docx import Document
from docx2pdf import convert

TEMPLATE = "SZABLON.docx"
DATA_FILE = "DANE.txt"
OUTPUT_DIR = "OUTPUT"

os.makedirs(OUTPUT_DIR, exist_ok=True)


# ---------- LOADER ----------
def loader(current, total, width=30):
    percent = current / total
    filled = int(width * percent)
    bar = "#" * filled + "-" * (width - filled)
    sys.stdout.write(
        f"\rPrzetwarzanie: [{bar}] {current}/{total} ({int(percent*100)}%)"
    )
    sys.stdout.flush()
    if current == total:
        print()


# ---------- DANE ----------
def load_and_group_data():
    with open(DATA_FILE, encoding="utf-8") as f:
        lines = [l.strip() for l in f if l.strip()]

    headers = [h.strip("[]") for h in lines[0].split()]
    groups = defaultdict(list)

    for line in lines[1:]:
        values = shlex.split(line)
        record = dict(zip(headers, values))
        groups[record["ID_ZAWIAD"]].append(record)

    return groups


# ---------- DOCX ----------
def replace_placeholders(doc, mapping):
    for p in doc.paragraphs:
        for key, val in mapping.items():
            if key in p.text:
                p.text = p.text.replace(key, val)


def generate_docx(id_zawiad, records):
    doc = Document(TEMPLATE)
    base = records[0]

    # nagłówki
    replace_placeholders(doc, {
        "[ID_ZAWIAD]": id_zawiad,
        "[PODMIOT]": base["PODMIOT"],
        "[ADRES]": base["ADRES"],
        "[OBREB]": base["OBREB"],
    })

    # znajdź linię wzorcową
    insert_index = None
    for i, p in enumerate(doc.paragraphs):
        if "[DZIALKA]" in p.text:
            insert_index = i
            p.clear()
            break

    # wstaw linie działek
    for r in records:
        line = (
            f"Dz. ewid.: {r['DZIALKA']}  "
            f"godz.: {r['GODZINA']}"
        )
        doc.paragraphs[insert_index]._p.addnext(
            doc.add_paragraph(line)._p
        )
        insert_index += 1

    docx_path = os.path.join(OUTPUT_DIR, f"ZAWIADOMIENIE_{id_zawiad}.docx")
    pdf_path = docx_path.replace(".docx", ".pdf")

    doc.save(docx_path)
    convert(docx_path, pdf_path)


# ---------- MAIN ----------
def main():
    groups = load_and_group_data()
    total = len(groups)

    for idx, (id_zawiad, records) in enumerate(groups.items(), start=1):
        generate_docx(id_zawiad, records)
        loader(idx, total)

    print("Zakończono – wygenerowano wszystkie zawiadomienia.")


if __name__ == "__main__":
    main()
