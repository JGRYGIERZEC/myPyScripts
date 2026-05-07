import os
import subprocess
from pathlib import Path

import pandas as pd
from docx import Document
from tqdm import tqdm
from PyPDF2 import PdfMerger

try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False

TEMPLATE_DOCX = "SZABLON.docx"
EXCEL_FILE = "PODMIOTY.xlsx"
OUT_DIR = Path("OUT")
TMP_DIR = Path("TMP_DOCX")

# Map placeholders in the docx to column names in the Excel
PLACEHOLDERS = {
    "[id_zawiad]": "id_zawiad",
    "[nazwa]": "nazwa",
    "[ulica]": "ulica",
    "[nra]": "nra",
    "[nrl]": "nrl",
    "[kod]": "kod",
    "[miasto]": "miasto",
}

def ensure_dirs():
    OUT_DIR.mkdir(exist_ok=True)
    TMP_DIR.mkdir(exist_ok=True)

def read_excel(path):
    df = pd.read_excel(path, dtype=str)
    df = df.fillna("")
    return df

def replace_placeholders_in_docx(doc: Document, mapping: dict):
    for p in doc.paragraphs:
        for ph, val in mapping.items():
            if ph in p.text:
                inline = p.runs
                new_text = p.text.replace(ph, val)
                for i in range(len(inline)-1, -1, -1):
                    p._element.remove(inline[i]._element)
                p.add_run(new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for ph, val in mapping.items():
                        if ph in p.text:
                            inline = p.runs
                            new_text = p.text.replace(ph, val)
                            for i in range(len(inline)-1, -1, -1):
                                p._element.remove(inline[i]._element)
                            p.add_run(new_text)

def create_docx_from_template(template_path, out_docx_path, mapping):
    doc = Document(template_path)
    replace_placeholders_in_docx(doc, mapping)
    doc.save(out_docx_path)

def convert_docx_to_pdf(docx_path: Path, pdf_path: Path):
    if DOCX2PDF_AVAILABLE:
        try:
            docx2pdf_convert(str(docx_path), str(pdf_path))
            return True
        except Exception:
            pass
    try:
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf", str(docx_path),
            "--outdir", str(pdf_path.parent)
        ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        return pdf_path.exists()
    except Exception:
        return False

def format_address_fields(row):
    # Prepare nra and nrl with slash only if nrl is non-empty
    nra = str(row.get("nra", "")).strip()
    nrl = str(row.get("nrl", "")).strip()
    if nrl:
        # add space before slash for readability: "nra /nrl" if nra present, else "/nrl"
        if nra:
            combined_nra_nrl = f"{nra} /{nrl}"
        else:
            combined_nra_nrl = f"/{nrl}"
    else:
        combined_nra_nrl = nra
    return combined_nra_nrl

def main():
    ensure_dirs()
    df = read_excel(EXCEL_FILE)
    generated_pdfs = []

    for idx, row in tqdm(df.iterrows(), total=len(df), desc="Generowanie kopert"):
        # Build mapping for placeholders
        mapping = {}
        for ph, col in PLACEHOLDERS.items():
            mapping[ph] = str(row.get(col, "")) if col in row.index else ""

        # Special handling for nra and nrl: combine with slash if nrl present
        combined_nra_nrl = format_address_fields(row)
        # Replace placeholders [nra] and [nrl] on the envelope:
        mapping["[nra]"] = combined_nra_nrl if combined_nra_nrl else mapping.get("[nra]", "")
        mapping["[nrl]"] = ""  # clear [nrl] because we've merged it into [nra] display

        id_zawiad = mapping.get("[id_zawiad]", "").strip()
        if not id_zawiad:
            id_zawiad = f"no_id_{idx}"

        safe_name = "".join(c for c in id_zawiad if c.isalnum() or c in "-_.")
        out_docx = TMP_DIR / f"{safe_name}.docx"
        out_pdf = OUT_DIR / f"{safe_name}.pdf"

        create_docx_from_template(TEMPLATE_DOCX, out_docx, mapping)

        ok = convert_docx_to_pdf(out_docx, out_pdf)
        if not ok:
            print(f"Konwersja nie powiodła się dla {out_docx}. Sprawdź instalację MS Word lub LibreOffice.")
            continue

        generated_pdfs.append(out_pdf)

    if generated_pdfs:
        merger = PdfMerger()
        for pdf in tqdm(generated_pdfs, desc="Łączenie PDF"):
            merger.append(str(pdf))
        merged_path = OUT_DIR / "merged.pdf"
        with open(merged_path, "wb") as f_out:
            merger.write(f_out)
        merger.close()
        print(f"Wygenerowano {len(generated_pdfs)} plików PDF. Połączono do: {merged_path}")
    else:
        print("Brak wygenerowanych PDF-ów.")

if __name__ == "__main__":
    main()
